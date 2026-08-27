from Classes.Documento import Documento
from Services.CNPJ import buscar_cnpj

from docx import Document

import locale 

class Apostilamento(Documento):

    def __init__(self, processo, secretaria, modelo, destino):
        self.processo = processo 
        self.secretaria = secretaria
        self.modelo = modelo
        self.destino = destino

    
    def set_mapeamento(self, sheet): 
        mapeamento = {
            "razao_social": {"iniciar": sheet["B2"].value},
            "troca_endereco": {"iniciar": sheet["B3"].value, "endereco_antigo": ""},
            "troca_marca": {"iniciar": sheet["B5"].value, "itens": ["***"]},
            "remanejamento": {"iniciar": sheet["B9"].value, "origem": ["***"], "destino": ["***"]}
        }
        self.mapeamento = mapeamento
    
    def criar_documento(self):

        processo = self.processo 
        secretaria = self.secretaria 
        
        doc = Document(self.modelo)
        
        def definir_tipo():
            if ("registro de preços" in self.processo.objeto.lower()):
                return "ATA DE REGISTRO DE PREÇOS"
            return "CONTRATO ADMINISTRATIVO"
        
        def definir_cabecalho(): 
            
            tipo = definir_tipo()
            cabecalho = f"APOSTILAMENTO N° XX - {tipo} N. {processo.n_contrato}\nProcesso Administrativo n° {processo.n_processo}\n{processo.modalidade} {processo.n_modalidade}\nEmpresa: {processo.fornecedor}\nOBJETO DO PROCESSO: {processo.objeto}."
            Documento.criar_texto(doc.add_paragraph(), cabecalho, negrito = False, posicionamento="Centro", fonte = "Arial")
            pass 
        
        def definir_objeto_apostilamento():
            texto_um = "OBJETO DO APOSTILAMENTO: Alteração d"
            
            #dados = buscar_cnpj(processo.cnpj)
            
            if (self.mapeamento["razao_social"]["iniciar"]): 
                texto_um += "a Razão Social "
                texto_dois = f"\nPassando de {processo.fornecedor} para RK MAGAZINE LTDA"
            if (self.mapeamento["troca_endereco"]["iniciar"]): 
                texto += "e endereço "
            if (self.mapeamento["troca_marca"]["iniciar"]): 
                texto += "e marcas "
            if (self.mapeamento["troca_marca"]["iniciar"]): 
                texto += "e estrutura orçamentária "

            texto = texto_um + texto_dois
            
            Documento.criar_texto(doc.add_paragraph(), texto, negrito = False, posicionamento="Esquerda", fonte = "Arial")

        def definir_restante_texto():
            texto = "Considerando que a alteração da razão social e endereço da contratada não implica modificação da personalidade jurídica da empresa, permanecendo inalterado o número de inscrição no CNPJ e as condições de habilitação exigidas, nos termos do art. 68 da Lei nº 14.133/2021;\n Considerando que a presente alteração possui natureza meramente cadastral, sem modificação do objeto contratado, quantitativos, valores ou condições de execução;\nConsiderando os princípios da legalidade, eficiência, continuidade do serviço público e interesse público previstos na Lei nº 14.133/2021\n;Fica autorizada a formalização da presente alteração mediante apostilamento"
            Documento.criar_texto(doc.add_paragraph(), texto, negrito = False, posicionamento="Esquerda", fonte = "Arial")

        def definir_data():
            texto = f"Bataguassu/MS, {Documento.encontrar_data_de_hoje_em_extenso()}."
            Documento.criar_texto(doc.add_paragraph(), texto, negrito = False, posicionamento="Direita", fonte = "Arial")

        def adicionar_assinatura():
            Documento.adicionar_linha_de_assinatura(doc.add_paragraph(), f"{secretaria.fundo}\n{secretaria.nome}")
            
        definir_cabecalho()
        definir_objeto_apostilamento()
        definir_restante_texto()
        definir_data()
        adicionar_assinatura()
        
        doc.save(f"{self.destino}\APOSTILAMENTO N° XX.2026 - {processo.fornecedor}.docx")

     
    """
    def __init__(self, contratado, liquidacao, valor, data, modelo, protocolo, endereco):
        super().__init__(contratado)
        self.liquidacao = liquidacao
        self.valor = valor
        self.data = data 
        self.modelo = modelo
        self.protocolo = protocolo
        self.endereco = endereco
    def criar_arquivo(self, termos):
    
        
        doc = Document(self.endereco)
        
        def formatar_data(self, objeto):
            return objeto.date().strftime("%d/%m/%Y")
        
        def converter_currency(self, valor):
            locale.setlocale(locale.LC_ALL, "pt_BR.UTF-8")
            return locale.currency(float(valor), grouping =True)
                    
        def adicionar_protocolo(self):
            substituir_protocolo = doc.paragraphs[1]
            substituir_protocolo.text = ""
            Documento.criar_texto(substituir_protocolo, f"PROTOCOLO DE RECEBIMENTO - NÚMERO {self.protocolo}", negrito = True)
               
        def criar_tabela(self, termos):
            
            for i in range(len(termos)):
                
                tabela = doc.tables[0]   
                nova_linha = tabela.add_row()
                                
                coluna_um = nova_linha.cells[0].paragraphs[0]
                coluna_dois = nova_linha.cells[1].paragraphs[0]
                coluna_tres = nova_linha.cells[2].paragraphs[0]
                coluna_quatro = nova_linha.cells[3].paragraphs[0]
                                    
                Documento.criar_texto(coluna_um, termos[i].contratado,  px = 8, negrito = True, fonte = "Arial")
                Documento.criar_texto(coluna_dois, termos[i].liquidacao,  px = 8, negrito = True, fonte = "Arial")
                Documento.criar_texto(coluna_tres, formatar_data(self, termos[i].data),  px = 8, negrito = True, fonte = "Arial")
                Documento.criar_texto(coluna_quatro, converter_currency(self, termos[i].valor),  px = 8, negrito = True, fonte = "Arial")
                                
        adicionar_protocolo(self) 
        criar_tabela(self, termos)
        
        doc.save(self.endereco)
    """