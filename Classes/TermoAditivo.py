from docx import Document

from Services.Data import EncontrarData
from Services.Number import converter_currency, pegar_numero_extenso

from Classes.Documento import Documento 


class TermoAditivo:
    
    cont = 0 
    
    def __init__(self, processo, secretaria, modelo, destino):
        self.processo = processo 
        self.secretaria = secretaria
        self.modelo = modelo
        self.destino = destino
        self.fundamento = "O presente Termo Aditivo tem fundamento legal no"
        self.prazo = False
        self.reajuste = 'Petal'
        
    def set_mapeamento(self, sheet):
        self.mapeamento = {
            "numero_termo": sheet["B2"].value,
            "valor": sheet["B3"].value,
            "parametros": {
                "prazo": {"iniciar": sheet["B4"].value, "inicio": sheet["C5"].value, "fim": sheet["C6"].value},
            }
        }
        
        self.numero_termo = self.mapeamento["numero_termo"]
        self.valor = self.mapeamento["valor"]
        
    def set_prazo(self, inicio, fim, tempo):
        self.inicio = inicio 
        self.fim = fim 
        self.tempo = tempo
        self.fundamento += "art. 84, da Lei n° 14.133/2021, e artigos 22 e 23, do Decreto Municipal nº 72/2025."
        
    def criar_documento(self):
        
        doc = Document(self.modelo)
        
        secretaria = self.secretaria 
        processo = self.processo 
        
        CLAUSULAS = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X"]

    
        
        def definir_parametros():
            if ((self.mapeamento["parametros"]["prazo"]["iniciar"])):
                
                self.prazo = True 
                
                dif = ((self.mapeamento["parametros"]["prazo"]["fim"] - self.mapeamento["parametros"]["prazo"]["inicio"]).days) // 30
                inicio = f"{EncontrarData("dia", self.mapeamento["parametros"]["prazo"]["inicio"].strftime("%d/%m/20%y"))} de {EncontrarData("mes", self.mapeamento["parametros"]["prazo"]["inicio"].strftime("%d/%m/20%y"), True).lower()} de {EncontrarData("ano", self.mapeamento["parametros"]["prazo"]["inicio"].strftime("%d/%m/20%y"))}"
                fim = f"{EncontrarData("dia", self.mapeamento["parametros"]["prazo"]["fim"].strftime("%d/%m/20%y"))} de {EncontrarData("mes", self.mapeamento["parametros"]["prazo"]["fim"].strftime("%d/%m/20%y"), True).lower()} de {EncontrarData("ano", self.mapeamento["parametros"]["prazo"]["fim"].strftime("%d/%m/20%y"))}"
                
                self.set_prazo(inicio, fim, dif)
        
        def criar_cabecalho():
            def definir_tipo(obj): 
                if ("registro de preços" in obj.lower()): 
                    return "A ATA DE REGISTRO DE PREÇOS N° "
                return "AO CONTRATO ADMINISTRATIVO N° "
            texto = f"{self.numero_termo}° TERMO ADITIVO {definir_tipo(processo.objeto)}{processo.n_contrato}"
            Documento.criar_texto(doc.add_paragraph(), texto, posicionamento="Centro", negrito = True, fonte = "Arial")
            texto = f"OBJETO DO PROCESSO: {processo.objeto}"
            Documento.criar_texto(doc.add_paragraph(), texto, posicionamento="Esquerda", negrito = True, fonte = "Arial")

        def criar_preambulo():
            texto = (f"{CLAUSULAS[self.cont]}. DAS PARTES: O {secretaria.fundo.upper()}/MS, pessoa jurídica de direito público interno, inscrito(a) no CNPJ sob o nº {secretaria.cnpj}, " 
                    f"situado a {secretaria.endereco}, Município de Bataguassu/MS, neste ato representado, conforme Decreto nº {secretaria.decreto} de {secretaria.data}, " 
                    f"pela Sra. {secretaria.nome}, {secretaria.cargo}, portadora da Matrícula Funcional nº {secretaria.matricula}, doravante denominada CONTRATANTE")
            if (secretaria.fundo.upper() == "MUNICIPIO DE BATAGUASSU"):
                texto = f"O {secretaria.fundo.upper()}/MS, pessoa jurídica de direito público interno, com sede à Avenida Aquidauana, nº 1001, Centro, Bataguassu/MS inscrito(a) no CNPJ sob o nº {secretaria.cnpj}, neste ato representado, por {secretaria.nome.upper()}, {secretaria.cargo.upper()}, portadora da Matrícula Funcional nº {secretaria.matricula}, doravante denominada CONTRATANTE"
            texto += f"e a empresa {processo.fornecedor}, inscrita no CNPJ n° {processo.cnpj}, com sede à SEI LA na cidade de SEI LÁ, neste ato representada pelo Sr. tal, doravante designadas CONTRATADAS, resolvem celebrar o presente Termo Aditivo, mediante as cláusulas e condições a seguir estabelecidas"
            self.cont += 1
            Documento.criar_texto(doc.add_paragraph(), texto, negrito = False, posicionamento="Esquerda", fonte = "Arial")
                   
        def criar_autorizacao():
            texto = f"{CLAUSULAS[self.cont]}. DA AUTORIZAÇÃO: O presente Termo Aditivo é celebrado em decorrência da autorização da Sra. {secretaria.cargo}, exarada em despacho constante do PROCESSO Nº {processo.n_processo}, {processo.modalidade} Nº {processo.n_modalidade}."
            self.cont += 1
            Documento.criar_texto(doc.add_paragraph(), texto, negrito = False, posicionamento="Esquerda", fonte = "Arial")
                    
        def criar_fundamento_legal():
            texto = f"{CLAUSULAS[self.cont]}. FUNDAMENTO LEGAL: {self.fundamento}"
            self.cont += 1
            Documento.criar_texto(doc.add_paragraph(), texto, negrito = False, posicionamento="Esquerda", fonte = "Arial")
            pass 
        
        def criar_objeto():
            texto = f"{CLAUSULAS[self.cont]}. DO OBJETO: Constitui objeto deste instrumento"
            if (self.prazo):
                texto += f" a prorrogação do prazo de vigência da ata por {self.tempo} ({(pegar_numero_extenso(self.tempo))}) meses e renovação proporcional dos quantitativos ao período, conforme ofícios, justificativas e parecer jurídico acostado aos autos."
            self.cont += 1
            Documento.criar_texto(doc.add_paragraph(), texto, negrito = False, posicionamento="Esquerda", fonte = "Arial")
        
        def criar_clausula_valor():
            if (self.reajuste):
                pass 
            texto = f"{CLAUSULAS[self.cont]}. DO VALOR: O valor do presente aditivo será de R$ {converter_currency(self.valor)} ({pegar_numero_extenso(self.valor, cambio = True)})."
            self.cont += 1
            Documento.criar_texto(doc.add_paragraph(), texto, negrito = False, posicionamento="Esquerda", fonte = "Arial")
            
        def criar_clausula_prazo():
            texto = f"{CLAUSULAS[self.cont]}. DO PRAZO: Fica prorrogado por mais {self.tempo} ({pegar_numero_extenso(self.tempo)}) meses com início em {self.inicio} e término em {self.fim}."
            self.cont += 1
            Documento.criar_texto(doc.add_paragraph(), texto, negrito = False, posicionamento="Esquerda", fonte = "Arial")

        def criar_clausulas_restantes():
            texto = f"{CLAUSULAS[self.cont]}. DA RATIFICAÇÃO: Ficam mantidas e ratificadas as demais cláusulas e condições da ata originária, naquilo que não contraria o presente termo aditivo.\n{CLAUSULAS[self.cont]}. DA PUBLICAÇÃO: Incumbirá ao contratante divulgar o presente instrumento no Portal Nacional de Contratações Públicas (PNCP), na forma prevista no art. 94 da Lei nº 14.133, de 2021 ou no Diário Oficial do órgão, bem como no respectivo sítio oficial na Internet, em atenção ao art. 91, caput, da Lei nº 14.133, de 2021, e ao art. 8º, §2º, da Lei nº 12.527, de 2011, c/c art. 7º, §3º, inciso V, do Decreto nº 7.724, de 2012."
            Documento.criar_texto(doc.add_paragraph(), texto, negrito = False, posicionamento="Esquerda", fonte = "Arial")
            Documento.criar_texto(doc.add_paragraph(), "Bataguassu/MS, data de assinatura digital.", negrito = False, posicionamento="Direita", fonte = "Arial")
        
        def criar_assinaturas():
            Documento.adicionar_linha_de_assinatura(doc.add_paragraph(), f"{secretaria.nome.upper()}\n{secretaria.cargo.upper()}")
            Documento.adicionar_linha_de_assinatura(doc.add_paragraph(), f"{processo.fornecedor}\nCNPJ N° {processo.cnpj}")
            pass 
    
        definir_parametros()
        Documento.adicionar_espaco(doc, 3)
        criar_cabecalho()
        Documento.adicionar_espaco(doc, 3)
        criar_preambulo()
        Documento.adicionar_espaco(doc, 3)
        criar_autorizacao()
        Documento.adicionar_espaco(doc, 3)
        criar_fundamento_legal()
        Documento.adicionar_espaco(doc, 3)
        criar_objeto()
        Documento.adicionar_espaco(doc, 3)
        criar_clausula_valor()
        Documento.adicionar_espaco(doc, 3)

        # criar clausula para itens?
        if (self.prazo): criar_clausula_prazo()
        if (self.reajuste): pass 
        
        Documento.adicionar_espaco(doc, 3)
        criar_clausulas_restantes()
        Documento.adicionar_espaco(doc, 3)
        criar_assinaturas()
        
        doc.save(f"{self.destino}/TERMO ADITIVO - MODELO.docx")
