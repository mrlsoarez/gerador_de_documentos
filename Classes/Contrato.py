from Entities.Processo import Processo
from Entities.Item import Item 

from docx import Document

from Services.CNPJ import buscar_cnpj
from Services.Number import converter_currency, pegar_numero_extenso

from Classes.Documento import Documento

class Contrato:
    def __init__(self, processo, secretaria, modelo, destino):
        self.processo = None 
        self.secretaria = secretaria
        self.modelo = rf"C:\Users\Usuario\Documents\MRL\MODELOS BASE\3. MODELOS DE DOCUMENTO\MODELO CONTRATO.docx"
        self.destino = destino
        
    def set_mapeamento(self, sheet):
        
        def definir_itens():
            contador = 10
            array = []
            while True: 
                if (sheet["A" + str(contador)].value == None):
                    break 
                array.append(
                    Item(
                        sheet["A" + str(contador)].value,
                        sheet["B" + str(contador)].value,
                        sheet["C" + str(contador)].value,
                        sheet["D" + str(contador)].value,
                        sheet["E" + str(contador)].value,
                        sheet["F" + str(contador)].value,
                        sheet["G" + str(contador)].value
                    )
                )
                
                contador += 1
            return array
        
        self.mapeamento = {
            "n_contrato": sheet["B2"].value,
            "processo": sheet["B3"].value,
            "cnpj": sheet["B4"].value,
            "modalidade": sheet["B5"].value,
            "n_modalidade": sheet["B6"].value, 
            "objeto": sheet["B7"].value,
        }
        
        self.set_info_processo()   
        self.itens = definir_itens()   
        self.set_valor_total()
        
    def set_info_processo(self):
        self.processo = Processo(self.mapeamento["processo"], self.mapeamento["cnpj"], "", self.mapeamento["modalidade"], self.mapeamento["n_modalidade"], self.mapeamento["n_contrato"], self.mapeamento["objeto"])

    def set_info_contratada(self):
        
        dados = None 
        
        try:
            dados = buscar_cnpj(self.processo.cnpj)
        except:
            print("Algo deu errado!")
        else:
            self.contratada_nome = dados["razao_social"]
            self.contratada_nome_rua = dados["estabelecimento"]["logradouro"]
            self.contratada_tipo_log = dados["estabelecimento"]["tipo_logradouro"]
            self.contratada_nome_bairro = dados["estabelecimento"]["bairro"]
            self.contratada_nome_cidade = dados["estabelecimento"]["cidade"]["nome"]
            self.contratada_nome_estado = dados["estabelecimento"]["estado"]["sigla"]
            self.contratada_numero = dados["estabelecimento"]["numero"]
            self.contratada_email = dados["estabelecimento"]["email"]
    
    def set_valor_total(self):
        cont = 0
        for i in range(len(self.itens)):
            cont += int(self.itens[i].v_tot)
        self.valor_total = cont 
            
    def buscar_info_empresa(self): 
        pass 
    
    def criar_documento(self):
        
        doc = Document(self.modelo)

        processo = self.processo 
        secretaria = self.secretaria
        
        #self.set_info_contratada()
        
        def definir_cabecalho(): 
            texto = f"CONTRATO ADMINISTRATIVO Nº {processo.n_contrato}\nPROCESSO ADMINISTRATIVO Nº {processo.n_processo}\n{processo.modalidade} Nº {processo.n_modalidade}"
            Documento.criar_texto(doc.add_paragraph(), texto, negrito = True, posicionamento="Centro", fonte = "Arial")
            texto2 = f"CONTRATO ADMINISTRATIVO Nº {processo.n_contrato} QUE FAZEM ENTRE SI O {secretaria.fundo.upper()} DE BATAGUASSU/MS E A EMPRESA {self.contratada_nome}."
            Documento.criar_texto(doc.add_paragraph(), texto2, negrito = True, posicionamento="Direita", fonte = "Arial")

        def definir_texto_inicial_contratante():
            texto = (f"O {secretaria.fundo.upper()}/MS, pessoa jurídica de direito público interno, inscrito(a) no CNPJ sob o nº {secretaria.cnpj}, " 
                     f"situado a {secretaria.endereco}, Município de Bataguassu/MS, neste ato representado, conforme Decreto nº {secretaria.decreto} de {secretaria.data}, " 
                     f"pela Sra. {secretaria.nome}, {secretaria.cargo}, portadora da Matrícula Funcional nº {secretaria.matricula}, doravante denominada CONTRATANTE")
            if (secretaria.fundo.upper() == "MUNICIPIO DE BATAGUASSU"):
                texto = f"O {secretaria.fundo.upper()}/MS, pessoa jurídica de direito público interno, com sede à Avenida Aquidauana, nº 1001, Centro, Bataguassu/MS inscrito(a) no CNPJ sob o nº {secretaria.cnpj}, neste ato representado, por {secretaria.nome.upper()}, {secretaria.cargo.upper()}, portadora da Matrícula Funcional nº {secretaria.matricula}, doravante denominada CONTRATANTE"
            Documento.criar_texto(doc.add_paragraph(), texto, negrito = False, posicionamento="Esquerda", fonte = "Arial")
 
        def definir_texto_inicial_contratada():
            texto = f"e a empresa {self.contratada_nome}, inscrita no CNPJ/MF sob o n° {processo.cnpj}, endereço eletrônico: {self.contratada_email}, sediada na {self.contratada_tipo_log} {self.contratada_nome_rua}, n° {self.contratada_numero}, Bairro: {self.contratada_nome_bairro}, na cidade de {self.contratada_nome_cidade}/{self.contratada_nome_estado}, doravante designada CONTRATADA, neste ato representada por [XXXXX], [FUNÇÃO], conforme atos constitutivos da empresa, tendo em vista o que consta no Processo Administrativo nº {processo.n_processo} e em observância às disposições da Lei nº 14.133, de 1º de abril de 2021, e demais legislação aplicável, resolvem celebrar o presente Termo de Contrato mediante as cláusulas e condições a seguir enunciadas."
            Documento.criar_texto(doc.add_paragraph(), texto, negrito = False, posicionamento="Esquerda", fonte = "Arial")

        def definir_itens():
            
            tabela = doc.tables[0]
    
            for i in range(len(self.itens)):
                
                nova_linha = tabela.add_row()
                                                            
                coluna_um = nova_linha.cells[0].paragraphs[0]
                coluna_dois = nova_linha.cells[1].paragraphs[0]
                coluna_tres = nova_linha.cells[2].paragraphs[0]
                coluna_quatro = nova_linha.cells[3].paragraphs[0]
                coluna_cinco = nova_linha.cells[4].paragraphs[0]
                coluna_seis = nova_linha.cells[5].paragraphs[0]
                coluna_sete = nova_linha.cells[6].paragraphs[0]
                
                item = self.itens[i] 
                
                Documento.criar_texto(coluna_um, item.n_item,  px = 11, fonte = "Arial")
                Documento.criar_texto(coluna_dois, item.cod,  px = 11,  fonte = "Arial")
                Documento.criar_texto(coluna_tres, item.descricao,  px = 11,  fonte = "Arial")
                Documento.criar_texto(coluna_quatro, item.unidade,  px = 11, fonte = "Arial")
                Documento.criar_texto(coluna_cinco, item.quant,  px = 11, fonte = "Arial")
                Documento.criar_texto(coluna_seis, converter_currency(item.v_unit),  px = 11, fonte = "Arial")
                Documento.criar_texto(coluna_sete, converter_currency(item.v_tot),  px = 11, fonte = "Arial")

            """
            
            tabela.add_row()
            dados = [
                {"celula": (7, 7), "conteudo": f"TOTAL: {converter_currency(self.valor_total)}", "negrito": True },
            ] 
            tabela = Documento.modificar_tabela(tabela, dados, fonte = "Arial")   
            """
             
        def definir_valor():
            texto = f"O valor total da contratação é de {converter_currency(self.valor_total)} ({pegar_numero_extenso(self.valor_total)})"
            Documento.criar_texto(doc.add_paragraph(), texto, negrito = False, posicionamento="Esquerda", fonte = "Arial")
            pass
        
        def adicionar_data():
            texto = "Bataguassu/MS, data de assinatura digital"
            Documento.criar_texto(doc.add_paragraph(), texto, negrito = False, posicionamento="Direita", fonte = "Arial")

        def adicionar_assinaturas():
            Documento.adicionar_linha_de_assinatura(doc.add_paragraph(), f"{secretaria.fundo.upper()}\n{secretaria.nome.upper()}")
            Documento.adicionar_linha_de_assinatura(doc.add_paragraph(), f"REPRESENTANTE\nCONTRATADA")
        
        def criar_extrato():
            doc2 = Document(self.modelo)
            Documento.criar_texto(doc2.add_paragraph(), f"EXTRATO CONTRATO ADMINISTRATIVO N° {processo.n_contrato} \nPROCESSO ADMINISTRATIVO N° {processo.n_processo} \n{processo.modalidade}N°{processo.n_modalidade}", negrito = True, posicionamento="Centro", fonte = "Arial")
            Documento.adicionar_espaco(doc, 4)
            Documento.criar_texto(doc2.add_paragraph(), f"PARTES: O {secretaria.fundo.upper()}/MS e a empresa exemplo", negrito = True, posicionamento="Esquerda", fonte = "Arial")
            Documento.adicionar_espaco(doc, 4)
            Documento.criar_texto(doc2.add_paragraph(), f"OBJETO: {processo.objeto}\n", negrito = True, posicionamento="Esquerda", fonte = "Arial")
            Documento.adicionar_espaco(doc, 4)            
            Documento.criar_texto(doc2.add_paragraph(), f"VIGÊNCIA: XXXXXXXXXXXXXXX\n", negrito = True, posicionamento="Esquerda", fonte = "Arial")
            Documento.adicionar_espaco(doc, 4)                
            Documento.criar_texto(doc2.add_paragraph(), f"PREÇO: O valor total da contratação é de {converter_currency(self.valor_total)} ({pegar_numero_extenso(self.valor_total)})\n", negrito = True, posicionamento="Esquerda", fonte = "Arial")
            Documento.adicionar_espaco(doc, 4)            
            Documento.criar_texto(doc2.add_paragraph(), f"DOTAÇÃO ORÇAMENTÁRIA: XXXXXXXXXXXXXXX\n", negrito = True, posicionamento="Esquerda", fonte = "Arial")
            Documento.adicionar_espaco(doc, 4)        
            Documento.criar_texto(doc2.add_paragraph(), f"DATA ASSINATURA: Bataguassu/MS, data de assinatura digital\n", negrito = True, posicionamento="Esquerda", fonte = "Arial")
            Documento.adicionar_espaco(doc, 4)    
            Documento.criar_texto(doc2.add_paragraph(), f"ASSINAM: O {secretaria.fundo.upper()}/MS e a empresa {processo.fornecedor}", negrito = True, posicionamento="Esquerda", fonte = "Arial")
            Documento.adicionar_espaco(doc, 4)    
            Documento.adicionar_linha_de_assinatura(doc2.add_paragraph(), "MURILO SOARES DE OLIVEIRA\nSETOR DE CONTRATOS")
            doc2.save(rf"{self.destino}/EXTRATO CONTRATO ADM.docx")
            
        definir_cabecalho()
        definir_texto_inicial_contratante()
        Documento.adicionar_espaco(doc, 5)
        definir_texto_inicial_contratada()
        Documento.adicionar_espaco(doc, 5)
        definir_itens()
        Documento.adicionar_espaco(doc, 5)
        definir_valor()
        Documento.adicionar_espaco(doc, 5)
        adicionar_data()
        Documento.adicionar_espaco(doc, 5)
        adicionar_assinaturas()
        
        criar_extrato()
        
        doc.save(rf"{self.destino}/CONTRATO ADM.docx")

   