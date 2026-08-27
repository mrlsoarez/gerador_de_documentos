from Classes.Documento import Documento
from Entities.Fiscal import Fiscal
from Services.Number import pegar_numero_extenso
from docx import Document

# Onde eu parei: definir o documento de portaria utilizando como base o modelo pronto
    # Já contém portaria e etc.
class Portaria(Documento):

    def __init__(self, processo, secretaria, modelo, destino):
        self.processo = processo 
        self.secretaria = secretaria
        self.modelo = r"C:\Users\Usuario\Documents\MRL\MODELOS BASE\3. MODELOS DE DOCUMENTO\MODELO DE PORTARIA.docx"
        self.destino = destino
        self.fiscais = None
    
    def set_mapeamento(self, sheet):
        
        def definir_fiscais():
            
            def alterar_info_estatutario(dado):
                if ("estatutário" in dado.lower() or "estatutario" in dado.lower()):
                    return "Efetivo"
                else:
                    return dado 
                
            contador = 6
            array = []
            while True: 
                if (sheet["B" + str(contador)].value == None):
                    break 
                array.append({
                    "principal": 
                        Fiscal(
                               sheet["B" + str(contador)].value,
                               sheet["C" + str(contador)].value,
                               sheet["D" + str(contador)].value,
                               alterar_info_estatutario(sheet["F" + str(contador)].value),
                               sheet["E" + str(contador)].value
                            ),
                    "suplente": 
                        Fiscal(
                            sheet["H" + str(contador)].value,
                            sheet["I" + str(contador)].value,
                            sheet["J" + str(contador)].value,
                            alterar_info_estatutario(sheet["L" + str(contador)].value),
                            sheet["K" + str(contador)].value
                        )
                })
                break
            return array
            
        mapeamento =  {
                    "tce": sheet["B2"].value,
                    "valor": sheet["B3"].value, 
                    "fiscais_planilha": definir_fiscais()
                }
        
        self.mapeamento = mapeamento
        self.fiscais = mapeamento["fiscais_planilha"]
        self.tce = mapeamento["tce"]
        self.valor = f"{mapeamento["valor"]} ({pegar_numero_extenso(mapeamento['valor'])})"
        
        

    def criar_documento(self):
        doc = Document(self.modelo)
                    
        def adicionar_titulo():
            doc.paragraphs[1].text = ""
            titulo = doc.paragraphs[1]
            texto = f"Portaria N° XXX - {Documento.encontrar_data_de_hoje_em_extenso()}"
            Documento.criar_texto(titulo, texto, negrito = True, posicionamento="Centro", fonte = "Arial")
        
        def adicionar_codigo_tce():
            doc.paragraphs[4].text = ""
            tce = doc.paragraphs[4]
            texto = f"Código Registro TCE: [{self.tce}]"
            Documento.criar_texto(tce, texto, negrito = True, posicionamento="Centro", fonte = "Arial")            
                      
        def criar_tabela_fiscais():
                            
            def capitalizar(texto):
                excecoes = {"de", "e"}
        
                return " ".join(
                    palavra if palavra.lower() in excecoes
                    else palavra.capitalize()
                    for palavra in texto.lower().split()
                )
                                
            def criar_nova_linha(dados):
                nova_linha = tabela.add_row()
                                
                coluna_um = nova_linha.cells[0].paragraphs[0]
                coluna_dois = nova_linha.cells[1].paragraphs[0]
                coluna_tres = nova_linha.cells[2].paragraphs[0]
                coluna_quatro = nova_linha.cells[3].paragraphs[0]
                                                        
                Documento.criar_texto(coluna_um, dados[0],  px = 11, fonte = "Arial")
                Documento.criar_texto(coluna_dois, dados[1],  px = 11,  fonte = "Arial")
                Documento.criar_texto(coluna_tres, dados[2],  px = 11,  fonte = "Arial")
                Documento.criar_texto(coluna_quatro, dados[3],  px = 11, fonte = "Arial")
                                
            fiscais = self.fiscais  
            tabela = doc.tables[0]   
                                                     
            for index in range(len(fiscais)):
                 
                criar_nova_linha(["", "FISCAL", "SUPLENTE", "GESTOR"])
                criar_nova_linha(["NOME DO SERVIDOR", capitalizar(fiscais[index]["principal"].nome), capitalizar(fiscais[index]["suplente"].nome), "Murilo Soares de Oliveira"])
                criar_nova_linha(["CARGO", capitalizar(fiscais[index]["principal"].cargo), capitalizar(fiscais[index]["suplente"].cargo), "Assistente de Administração"])
                criar_nova_linha(["MATRÍCULA", fiscais[index]["principal"].matricula, fiscais[index]["suplente"].matricula, "117810"])
                criar_nova_linha(["VÍNCULO", capitalizar(fiscais[index]["principal"].vinculo), capitalizar(fiscais[index]["suplente"].vinculo), "Efetivo"])
                criar_nova_linha(["SECRETARIA", capitalizar(fiscais[index]["principal"].secretaria), capitalizar(fiscais[index]["suplente"].secretaria), "Administração e Finanças"])
        
                if (index != len(fiscais) - 1): 
                    nova_linha = tabela.add_row()
            
        def alterar_tabela_atas():
                        
                tabela = doc.tables[1] 
                texto = ""
                
                fornecedores = self.processo.fornecedor
                processo = self.processo
                        
                for i in range(len(fornecedores)):
                    texto += f"ATA N° {fornecedores[i]['n_contrato']} - {fornecedores[i]['fornecedor']}, CNPJ N° {fornecedores[i]['cnpj']}\n"
                        
                dados = [
                    {"celula": (0, 1), "conteudo": texto, "negrito": True },
                    {"celula": (1, 1), "conteudo": processo.objeto, "negrito": True },
                    {"celula": (2, 1), "conteudo": "12 meses.", "negrito": True },
                    {"celula": (3, 1), "conteudo": f"R$ {self.valor}", "negrito": True },
                ]  
                                    
                doc.tables[1] = Documento.modificar_tabela(doc.tables[1], dados, "Arial")
                pass   
                    
        def inserir_data_rodape():
            index = None
            for i in range(len(doc.paragraphs)):
                if ("Gabinete da Prefeita Municipal de Bataguassu-MS, em [DATA]" in doc.paragraphs[i].text):
                    index = i                  
            doc.paragraphs[index].text = ""
            paragrafo = doc.paragraphs[index] 
            texto = f"Gabinete da Prefeita Municipal de Bataguassu-MS, em {Documento.encontrar_data_de_hoje_em_extenso()}"
            Documento.criar_texto(paragrafo, texto, negrito = True, fonte = "Arial")
                    
          
        
                    
                
        """
                    
        def adicionar_titulo(self):
            texto = f"Portaria N° XX - {Documento.encontrar_data_de_hoje_em_extenso()}"
            Documento.criar_texto(doc.add_paragraph(), texto, negrito = True, posicionamento="Centro", fonte = "Arial")
        
        def adicionar_codigo_tce(self):
            texto = f"Código Registro TCE: [{self.tce}]"
            Documento.criar_texto(doc.add_paragraph(), texto, negrito = True, posicionamento="Centro", fonte = "Arial")
        """
        adicionar_titulo()
        adicionar_codigo_tce()
        criar_tabela_fiscais()
        alterar_tabela_atas()
        inserir_data_rodape()

        doc.save(rf"{self.destino}\modelo.docx")
        
       
    """
        def __init__(self, numero_portaria, codigo_tce, fornecedores, objeto, valor_total, modelo, endereco):
            self.numero_portaria = numero_portaria
            self.codigo_tce = codigo_tce
            self.fornecedores = fornecedores
            self.objeto = objeto
            self.valor_total = valor_total
            self.fiscais = [] 
            self.modelo = modelo 
            self.endereco = endereco
        
        def criar_arquivo(self):
            
    """
            
               
